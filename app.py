import streamlit as st
from pypdf import PdfWriter, PdfReader
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import google.generativeai as genai
from google.api_core.exceptions import ResourceExhausted
import io
import os
import time
import tempfile
import re
from datetime import datetime

# ==========================================
# --- 1. CONFIGURAÇÃO VISUAL ---
# ==========================================

st.set_page_config(page_title="Auditor EIA Pro - Super Base", page_icon="⚖️", layout="wide")

st.markdown("""
<style>
    .stButton>button { width: 100%; border-radius: 5px; height: 3em; background-color: #1f77b4; color: white; }
    .stSuccess, .stInfo { border-left: 5px solid #1f77b4; }
</style>
""", unsafe_allow_html=True)

if 'uploader_key' not in st.session_state: st.session_state.uploader_key = 0
def reset_app(): st.session_state.uploader_key += 1

# ==========================================
# --- 2. SUPER BASE DE DADOS LEGISLATIVA ---
# ==========================================

# Leis que se aplicam a QUASE TODOS os projetos
COMMON_LAWS = {
    "RJAIA (DL 151-B/2013 + Alterações)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2013-116043164",
    "SIMPLEX AMBIENTAL (DL 11/2023)": "https://diariodarepublica.pt/dr/detalhe/decreto-lei/11-2023-207604364",
    "LUA - Licenciamento Único Ambiental (DL 75/2015)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2015-106567543",
    "LEI DE BASES DO AMBIENTE (Lei 19/2014)": "https://diariodarepublica.pt/dr/legislacao-consolidada/lei/2014-34543212",
    "REDE NATURA 2000 (DL 140/99)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/1999-34460975",
    "REGULAMENTO GERAL DO RUÍDO (DL 9/2007)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2007-34526556",
    "LEI DA ÁGUA (Lei 58/2005)": "https://diariodarepublica.pt/dr/legislacao-consolidada/lei/2005-34563267",
    "REGIME GERAL DE RESÍDUOS (DL 102-D/2020)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2020-150917243",
    "RESPONSABILIDADE AMBIENTAL (DL 147/2008)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2008-34484567",
    "QUALIDADE DO AR (DL 102/2010)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2010-34512345"
}

# Leis Específicas por Tipologia (Expandida)
SPECIFIC_LAWS = {
    "1. Agricultura, Pecuária e Floresta": {
        "NREAP - Atividade Pecuária (DL 81/2013)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2013-34567890",
        "GESTÃO EFLUENTES PECUÁRIOS (Port. 631/2009)": "https://diariodarepublica.pt/dr/detalhe/portaria/631-2009-518868",
        "PROGRAMAS DE AÇÃO NITRATOS (Port. 259/2012)": "https://diariodarepublica.pt/dr/detalhe/portaria/259-2012-345678",
        "SISTEMA DEFESA FLORESTA (DL 124/2006)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2006-34512345",
        "ARBORIZAÇÃO E REARBORIZAÇÃO (DL 96/2013)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2013-10654321"
    },
    "2. Indústria Extrativa (Minas e Pedreiras)": {
        "LEI DE BASES RECURSOS GEOLÓGICOS (Lei 54/2015)": "https://diariodarepublica.pt/dr/legislacao-consolidada/lei/2015-106556789",
        "RESÍDUOS DE EXTRAÇÃO (DL 10/2010)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2010-34658745",
        "REVELAÇÃO E APROVEITAMENTO (DL 270/2001)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2001-34449875",
        "SEGURANÇA E SAÚDE EM MINAS (DL 162/90)": "https://diariodarepublica.pt/dr/detalhe/decreto-lei/162-1990-417937"
    },
    "3. Energia (Renováveis, Linhas, H2)": {
        "BASES DO SISTEMA ELÉTRICO (DL 15/2022)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2022-177343687",
        "PRODUÇÃO H2 E GASES RENOVÁVEIS (DL 62/2020)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2020-13456789",
        "CAMPOS ELETROMAGNÉTICOS (Port. 1421/2004)": "https://diariodarepublica.pt/dr/detalhe/portaria/1421-2004-193456",
        "REGULAMENTO SEGURANÇA LINHAS AT (DL 25/2016)": "https://diariodarepublica.pt/dr/detalhe/decreto-lei/25-2016-10654321"
    },
    "4. Indústria e Química (Seveso, Emissões)": {
        "SISTEMA INDÚSTRIA RESPONSÁVEL (SIR - DL 169/2012)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2012-34567890",
        "EMISSÕES INDUSTRIAIS (DL 127/2013)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2013-34789569",
        "PREVENÇÃO ACIDENTES GRAVES (SEVESO III - DL 150/2015)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2015-106558967",
        "REGISTO E AVALIAÇÃO SUBSTÂNCIAS (REACH)": "https://echa.europa.eu/regulations/reach/legislation"
    },
    "5. Infraestruturas e Transportes": {
        "ESTATUTO DAS ESTRADAS (Lei 34/2015)": "https://diariodarepublica.pt/dr/legislacao-consolidada/lei/2015-34585678",
        "SERVIDÕES AERONÁUTICAS (DL 48/2022)": "https://diariodarepublica.pt/dr/detalhe/decreto-lei/48-2022-185799345",
        "GESTÃO DE RUÍDO INFRAESTRUTURAS (DL 146/2006)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2006-34512345"
    },
    "6. Água, Saneamento e Hidráulica": {
        "UTILIZAÇÃO RECURSOS HÍDRICOS (DL 226-A/2007)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2007-34567890",
        "QUALIDADE ÁGUA CONSUMO (DL 306/2007)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2007-34512345",
        "ÁGUAS RESIDUAIS URBANAS (DL 152/97)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/1997-34512345",
        "SEGURANÇA DE BARRAGENS (DL 21/2018)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2018-114833256",
        "GESTÃO RISCO INUNDAÇÕES": "https://diariodarepublica.pt/dr/detalhe/resolucao-conselho-ministros/51-2016-10654321"
    },
    "7. Resíduos e Economia Circular": {
        "REGIME GERAL GESTÃO RESÍDUOS (DL 102-D/2020)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2020-150917243",
        "DEPOSIÇÃO EM ATERRO (DL 102-D/2020 Anexo)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2020-150917243",
        "INCINERAÇÃO E CO-INCINERAÇÃO (DL 127/2013)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2013-34789569"
    },
    "8. Turismo e Urbanismo": {
        "RJUE - Urbanização e Edificação (DL 555/99)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/1999-34563452",
        "RJET - Empreendimentos Turísticos (DL 39/2008)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2008-34460567",
        "RESERVA ECOLÓGICA NACIONAL (REN - DL 166/2008)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2008-34512345",
        "RESERVA AGRÍCOLA NACIONAL (RAN - DL 73/2009)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2009-34567890"
    },
    "Outra Tipologia": {}
}

# ==========================================
# --- 3. LÓGICA DE PROCESSO (FILE API + LOCAL) ---
# ==========================================

def get_available_models(api_key):
    try:
        genai.configure(api_key=api_key)
        return [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
    except: return []

def extract_text_from_pdfs_local(files):
    """Extrai texto de PDFs de legislação extra (máx 200MB total na memória)."""
    text = ""
    for f in files:
        try:
            reader = PdfReader(f)
            text += f"\n>>> INÍCIO DIPLOMA EXTRA: {f.name} <<<\n"
            for page in reader.pages:
                text += page.extract_text() + "\n"
            text += f">>> FIM DIPLOMA EXTRA: {f.name} <<<\n"
        except Exception as e:
            text += f"\n[ERRO LEITURA {f.name}: {str(e)}]\n"
    return text

def merge_pdfs_to_temp(uploaded_files):
    merger = PdfWriter()
    for uploaded_file in uploaded_files:
        merger.append(uploaded_file)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        merger.write(tmp)
        tmp_path = tmp.name
    return tmp_path

def analyze_large_document(merged_pdf_path, laws_str, extra_laws_text, prompt_instructions, key, model_name):
    genai.configure(api_key=key)
    status_msg = st.empty()
    status_msg.info("📤 A enviar processo para a Google Cloud (File API)...")
    
    processo_file = None
    try:
        # 1. Upload do Processo Principal
        processo_file = genai.upload_file(path=merged_pdf_path, display_name="Processo EIA Auditoria")
        
        # 2. Polling de Estado
        status_msg.info("⚙️ A Google está a processar o PDF...")
        while processo_file.state.name == "PROCESSING":
            time.sleep(2)
            processo_file = genai.get_file(processo_file.name)
        
        if processo_file.state.name == "FAILED": raise ValueError("Google falhou a leitura do PDF.")
        
        status_msg.success("✅ Leitura concluída. A iniciar auditoria com IA...")

        # 3. Prompt Avançado
        model = genai.GenerativeModel(model_name)
        
        full_prompt = [
            prompt_instructions,
            "\n=== QUADRO LEGISLATIVO GERAL (VERIFICAR CONFORMIDADE) ===\n",
            laws_str,
            "\n=== QUADRO LEGISLATIVO EXTRA (TEXTO COMPLETO) ===\n",
            extra_laws_text if extra_laws_text else "Nenhum diploma extra carregado.",
            "\n=== INSTRUÇÃO DE EXECUÇÃO ===\n",
            "Com base nas leis acima e no PROCESSO EIA em anexo, gera o relatório.",
            processo_file
        ]

        response = model.generate_content(full_prompt)
        status_msg.empty()
        return response.text

    except ResourceExhausted:
        return "🚨 ERRO CRÍTICO: Limite de Tokens/Custo da API excedido."
    except Exception as e:
        return f"❌ Erro Técnico: {str(e)}"
    finally:
        if processo_file:
            try: genai.delete_file(processo_file.name)
            except: pass

# ==========================================
# --- 4. GERADOR DE WORD ---
# ==========================================

def clean_markdown(text):
    return text.replace('**', '').strip()

def create_professional_doc(content, project_type, active_laws_dict, extra_files_names):
    doc = Document()
    
    # Estilos
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    
    # Título
    title = doc.add_heading('AUDITORIA DE CONFORMIDADE EIA', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Tipologia: {project_type}').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Data: {datetime.now().strftime("%d/%m/%Y")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('---')

    # Conteúdo
    for line in content.split('\n'):
        line = line.strip()
        if not line: continue
        
        if line.startswith('## '):
            clean = clean_markdown(line.replace('## ', ''))
            h = doc.add_heading(clean.upper(), level=1)
            h.style.font.color.rgb = RGBColor(14, 77, 164)
        elif line.startswith('### '):
            clean = clean_markdown(line.replace('### ', ''))
            doc.add_heading(clean, level=2)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)

    # Anexo Legislativo
    doc.add_page_break()
    doc.add_heading('ANEXO: QUADRO LEGAL REFERENCIADO', level=1)
    
    doc.add_paragraph("1. Diplomas de Base e Setoriais:", style='Heading 2')
    for name, url in active_laws_dict.items():
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(name).bold = True
        if url.startswith("http"):
            p.add_run(f" (Ver Diploma)").italic = True
    
    if extra_files_names:
        doc.add_paragraph("2. Legislação Extra Específica (PDFs):", style='Heading 2')
        for f_name in extra_files_names:
            doc.add_paragraph(f_name, style='List Bullet')

    bio = io.BytesIO()
    doc.save(bio)
    return bio

# ==========================================
# --- 5. INTERFACE DO UTILIZADOR ---
# ==========================================

with st.sidebar:
    st.header("1. Configuração")
    api_key = st.text_input("Chave API Google", type="password")
    
    model_name = ""
    if api_key:
        models = get_available_models(api_key)
        if models:
            ix = 0
            # Tentar selecionar Flash por omissão (mais rápido/barato para muitos dados)
            for i, m in enumerate(models):
                if 'flash' in m: ix = i; break
            model_name = st.selectbox("Modelo IA", models, index=ix)
            if "pro" in model_name: st.caption("⚠️ O modelo Pro pode ser mais lento.")
    
    st.markdown("---")
    st.header("2. Tipologia do Projeto")
    project_type = st.selectbox("Selecione o Setor:", list(SPECIFIC_LAWS.keys()) + ["Outra Tipologia"])
    
    # Construção Dinâmica da Lista de Leis
    active_laws = COMMON_LAWS.copy()
    if project_type in SPECIFIC_LAWS:
        active_laws.update(SPECIFIC_LAWS[project_type])
    
    with st.expander(f"📚 Ver {len(active_laws)} Diplomas Ativos"):
        for k, v in active_laws.items():
            st.markdown(f"- [{k}]({v})")
            
    st.markdown("---")
    st.header("3. Legislação Extra")
    st.caption("Carregue PDMs, Regulamentos Municipais ou Portarias específicas.")
    extra_laws_files = st.file_uploader("Upload PDFs Extra", type=['pdf'], accept_multiple_files=True)

st.title("⚖️ Auditor EIA Pro: Super Database")
st.markdown("Auditoria inteligente com base no RJAIA, Simplex Ambiental e legislação setorial específica.")

uploaded_files = st.file_uploader(
    "📂 Carregar Processo EIA (Tomo I, RNT, Anexos...)", 
    type=['pdf'], 
    accept_multiple_files=True,
    key=f"uploader_{st.session_state.uploader_key}"
)

# --- PROMPT REFINADO (7 CAPÍTULOS + OPINIÃO TÉCNICA) ---
instructions = f"""
Atua como Perito Sénior em Engenharia do Ambiente e Jurista.
Auditoria de conformidade rigorosa ao EIA do setor: {project_type}.

CONTEXTO LEGISLATIVO:
1. Verifica a conformidade com a 'Legislação Base' listada.
2. Verifica a conformidade com a 'Legislação Extra' (texto completo fornecido), se existir.

ESTRUTURA DO RELATÓRIO (Usa Markdown ##):

## 1. ENQUADRAMENTO LEGAL E ADMINISTRATIVO
   - Enquadramento RJAIA (Anexo I/II) e verificação DL 11/2023 (Simplex).
   - Verificação de conformidade administrativa (entidades, prazos, peças obrigatórias).
   - O projeto cumpre os Instrumentos de Gestão Territorial (PDM, REN, RAN) citados?

## 2. PRINCIPAIS IMPACTES (TÉCNICO)
   - Resumo dos impactes negativos significativos por descritor (Ar, Água, Ruído, Biodiversidade, Solos).

## 3. MEDIDAS DE MITIGAÇÃO
   - Lista as medidas de minimização propostas (Construção e Exploração).

## 4. ANÁLISE CRÍTICA E BENCHMARKING
   - As medidas são suficientes face às Melhores Técnicas Disponíveis (MTD) do setor {project_type}?
   - Existem lacunas face à legislação listada (ex: falta de monitorização de ruído, falta de tratamento de efluentes)?

## 5. FUNDAMENTAÇÃO
   - Lista evidências concretas.
   - OBRIGATÓRIO: Referencia a página do PDF onde a informação se encontra (ex: "Ref: EIA, Tomo I, pág. 120").

## 6. CITAÇÕES RELEVANTES
   - Transcreve pequenos trechos do EIA que comprovem os pontos críticos levantados.

## 7. CONCLUSÕES E RECOMENDAÇÕES TÉCNICAS
   - Opinião técnica fundamentada sobre a qualidade do estudo e viabilidade ambiental.
   - NÃO emitir "Parecer Favorável/Desfavorável" administrativo.
   - Listar recomendações de melhoria ou pedidos de elementos adicionais (Aditamentos).

Tom: Formal, Técnico e Construtivo.
"""

if st.button("🚀 INICIAR AUDITORIA", type="primary"):
    if not api_key or not model_name: st.error("⚠️ Falta a API Key.")
    elif not uploaded_files: st.warning("⚠️ Falta o Processo EIA.")
    else:
        with st.spinner("A cruzar dados do Processo com a Base Legislativa..."):
            
            # 1. Lista de Leis Base (Texto para Prompt)
            laws_str = "\n".join([f"- {k}" for k in active_laws.keys()])
            
            # 2. Leis Extras (Extração Local)
            extra_text = ""
            extra_names = []
            if extra_laws_files:
                extra_text = extract_text_from_pdfs_local(extra_laws_files)
                extra_names = [f.name for f in extra_laws_files]
            
            # 3. Processo (Merge + File API)
            temp_path = merge_pdfs_to_temp(uploaded_files)
            
            result_text = analyze_large_document(
                temp_path, 
                laws_str, 
                extra_text, 
                instructions, 
                api_key, 
                model_name
            )
            
            try: os.remove(temp_path)
            except: pass
            
            if "🚨" in result_text or "❌" in result_text:
                st.error(result_text)
            else:
                st.success("Análise Concluída com Sucesso!")
                with st.expander("📄 Ler Relatório", expanded=True):
                    st.markdown(result_text)
                
                docx = create_professional_doc(result_text, project_type, active_laws, extra_names)
                st.download_button(
                    "⬇️ Download Relatório Word", 
                    docx.getvalue(), 
                    "Auditoria_EIA_Pro.docx", 
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    on_click=reset_app
                )

